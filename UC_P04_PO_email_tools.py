import os
from docx import Document
from datetime import date, datetime

OUT_DIR  = os.path.dirname(__file__)   # Generated RFQ files save here
vendor_emails = {
    "Hindalco Chemicals" : "sales@hindalcochemicals.com",
    "Indo Gulf Corp"     : "sales@indogulfcorp.com",
    "Gulf Chem Trading"  : "sales@gulfchemtrading.com",
    "Chengdu Carbon"     : "sales@chengducarbon.com",
    "Gulf Anode Co"      : "sales@gulfanodeco.com",
    "CPCL Suppliers"     : "sales@cpclsuppliers.com",
    "SRF Industries"     : "sales@srfindustries.com",
    "Bhushan Metals"     : "sales@bhushanmetals.com",
    "L&T Construction"   : "sales@l&tconstruction.com",
    "Adani Logistics"    : "sales@adanilogistics.com",
}

EMAILS = [
    {
        "id": 1,
        "subject": "Quotation for Caustic Soda",
        "sender": "sales@hindalcochemicals.com",
        "date": "2026-07-25",
        "attachments": [{"filename": "Quotation.docx"}, {"filename": "Banking.docx"}]
    },
    {
        "id": 2,
        "subject": "Meeting Reminder",
        "sender": "hr@example.com",
        "date": "2026-07-27",
        "attachments": []
    },
    {
        "id": 3,
        "subject": "Invoice",
        "sender": "finance@example.com",
        "date": "2026-07-29",
        "attachments": []
    },
        {
        "id": 4,
        "subject": "Quotation for Caustic Soda",
        "sender": "sales@indo.com",
        "date": "2026-07-24",
        "attachments": [{"filename": "Floatation.docx"}]
    }
]

def check_for_latest_quotation(item: str) -> str:
    '''
    Check for qoutation emails for a given item.
    If available, extract information from the attached document, present it to the user and ask them if the quote is ACCEPTABLE.
    If YES, generate a PO with the extracted information and email it to the vendor, ELSE email the vendor WITHOUT attachments to provide a better quote.
    '''
    keyword_emails = []

    # Check for emails with the given keyword along with the word 'quotation'
    # Only keep emails with attachments and authorized sender addresses (vendor_emails)
    for email in EMAILS:
        subject = email["subject"].lower()
        if item.lower() in subject and 'quotation' in subject and email["attachments"] and email["sender"] in vendor_emails.values():
            keyword_emails.append(email)


    # If no keyworded emails are found, return: No emails were found
    if not keyword_emails:
        return "No emails with attachments were found for the given keyword."

    # Get the latest of the emails containing the given keyword
    keyword_emails = max(keyword_emails, key=lambda email: datetime.strptime(email['date'], "%Y-%m-%d"))

    # Remove attachments that do not have the word 'quotation' and do not end with '.docx'
    keyword_emails["attachments"] = [attachment for attachment in keyword_emails["attachments"] if "quotation" in attachment["filename"].lower() and attachment["filename"].lower().endswith(".docx")]

    # If no attachments are left in the email, return: No relevant attachments found
    if not keyword_emails["attachments"]:
        return "No quotation docx attachments found in the email."
    

    # Determine attachment filepath, vendor email and vendor name
    filepath = os.path.join(OUT_DIR, keyword_emails["attachments"][0]["filename"])
    vendor_email = keyword_emails["sender"]
    vendor_name = next((k for k, v in vendor_emails.items() if v == vendor_email), None)

    # Check for tables in the attached document
    doc = Document(filepath)
    rows = []
    if not doc.tables:
        return "There are no tables in the document to parse."
    # Iterate through each row of the first available table
    row_count = 0
    for row in doc.tables[0].rows:
        # Extract text from each cell in the row
        row_data = [cell.text.strip() for cell in row.cells]
        if not row_count:
            row_data.insert(3, 'Vendor')
        else:
            row_data.insert(3, vendor_name)
        row_data = " | ".join(row_data)
        rows.append(row_data)
        row_count += 1

    table_string = "\n".join(rows)
    table_string = table_string + "\nVendor email: " + vendor_email
    return table_string

def generate_po(details: str) -> str:
    """
    Generate a formatted Word PO document from procurement details.
    Input format: 'item|quantity|unit|delivery_date|vendor|total'
    Example: 'Caustic Soda|500|MT|2024-07-15|Hindalco Chemicals|18250000'
    
    Use this tool AFTER retrieving the latest quotation details which the USER has ACCEPTED.
    """
    # Parse the pipe-delimited input string
    try:
        parts    = [p.strip() for p in details.split("|")]
        item     = parts[0] if len(parts) > 0 else "Item TBD"
        qty      = parts[1] if len(parts) > 1 else "TBD"
        unit     = parts[2] if len(parts) > 2 else "MT"
        del_date = parts[3] if len(parts) > 3 else str(date.today())
        vendor   = parts[4] if len(parts) > 4 else ["To be determined"]
        total    = parts[5] if len(parts) > 5 else "TBD"
    except Exception:
        item, qty, unit, del_date, vendor, total = "TBD", "TBD", "MT", str(date.today()), ["TBD"], "TBD"

    # Build the Word document using python-docx
    doc = Document()

    # Header section
    doc.add_heading("PURCHASE ORDER", level=0)
    doc.add_paragraph(f"PO No: NALCO/MAT/{date.today().strftime('%Y%m%d')}/AUTO-{item[:3].upper()}")
    doc.add_paragraph(f"Date: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph(f"To: {vendor}")
    doc.add_paragraph("")

    # Item details table
    doc.add_heading("Item Details", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Item / Material", "Quantity", "Unit", "Required Delivery Date", "Total Cost"]):
        hdr[i].text = h

    row = table.add_row().cells
    row[0].text = item
    row[1].text = str(qty)
    row[2].text = unit
    row[3].text = del_date
    row[4].text = total

    doc.add_paragraph("")

    # Terms and conditions
    doc.add_heading("Terms and Conditions", level=1)
    for tc in [
        "Delivery: As per schedule specified above. Delay penalty as per NALCO standard T&C.",
        "Payment: As per NALCO standard payment terms (45 days from delivery acceptance).",
        "Quality: Material must conform to NALCO technical specifications on file.",
        "Reference: This PO is governed by NALCO Standard Procurement Terms & Conditions.",
    ]:
        doc.add_paragraph(tc, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Please submit your queries to: procurement@nalco.gov.in by the date specified above.")
    doc.add_paragraph("For queries: Contact NALCO Materials Management Department.")

    # Save the document with a unique filename
    filename = f"PO_{item.replace(' ', '_')}_{date.today().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(OUT_DIR, filename)
    doc.save(filepath)

    return f"PO document generated and saved: {filename}\nEmail to: {vendor_emails[vendor]}"

def send_email(details: str) -> str:
    '''
    On confirmation from user, send an email to the specified vendor with a specified message.
    Input format: vendor_email | subject | message | attachment_name
    If attachments are not required, input attachment_name as None
    '''
    # Parse the pipe-delimited input string
    try:
        parts    = [p.strip() for p in details.split("|")]
        vendor_email    = parts[0]
        subject         = parts[1]
        message         = parts[2]
        attachment_name = parts[3]
    except Exception:
        return "Email couldn't be formatted correcty."

    attachment_format = []
    if attachment_name != "None":
        attachment_format = [{'filename': attachment_name}]

    email_format = {
        "to": vendor_email,
        "subject": subject,
        "message": message,
        "attachment": attachment_format
    }
    return f"Email has been sent to {email_format['to']} with subject: {email_format['subject']}\nMessage: {email_format['message']}\nAttachment: {email_format['attachment']}"

# Tests
if __name__ == "__main__":
    print(check_for_latest_quotation("caustic Soda"))
    print(generate_po("Caustic Soda | 500 | MT | 2026-07-29 | Hindalco Chemicals | 18250000"))
    print(send_email("sales@hindalcochemicals.com | Purchase order | PFA the required PO document | PO_Caustic_Soda_20260730.docx"))