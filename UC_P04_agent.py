"""
UC-P04 | Procurement AI Agent — Automated PO and RFQ Drafting
Local AI Agent (Ollama + SQLite + python-docx + Streamlit)
NALCO AI Training | Procurement Department

WHAT THIS APP BUILDS:
  An AI agent that accepts plain English procurement instructions and autonomously:
  1. Queries a local SQLite vendor database
  2. Retrieves historical purchase prices and market benchmarks
  3. Generates a formatted Word RFQ document
  4. Check for quotation emails
  5. Generate purchase order document

  The agent DECIDES which tools to call and in what order.
  You never hard-code the decision logic — the LLM does it.

WHAT IS AN AGENT (vs a chatbot)?
  Chatbot: User asks question -> LLM generates text answer
  Agent:   User gives goal -> LLM decides which TOOLS to call -> executes them -> returns result

HOW TO RUN:
  1. python UC_P04_setup_db.py   (creates procurement.db)
  2. streamlit run UC_P04_agent.py OR python -m streamlit run UC_P04_agent.py

PREREQUISITES:
  pip install streamlit ollama python-docx
  ollama pull qwen3.5
"""

# =============================================================================
# IMPORT BLOCK
# PROMPT TO REGENERATE:
# "Import os, BytesIO from io, date from datetime, connect from sqlite3,
#  streamlit, chat from ollama, and Document from python-docx.
#  Create a path DB_PATH to the SQLite database named "UC_P04_procurement.db".
#  Create a path OUT_DIR to the directory in which this file is saved."
# =============================================================================
import os
from io import BytesIO
from datetime import date

import UC_P04_PO_email_tools

from sqlite3 import connect
import streamlit as st
from ollama import chat
from docx import Document

# Path to the SQLite database created by UC_P04_setup_db.py
DB_PATH  = os.path.join(os.path.dirname(__file__), "UC_P04_procurement.db")
OUT_DIR  = os.path.dirname(__file__)   # Generated RFQ files save here

# =============================================================================
# TOOL 1: GET APPROVED VENDORS
# PROMPT TO REGENERATE:
# "Write a tool function called get_approved_vendors(category: str)
#  that queries a SQLite table 'vendors' for rows where category matches the
#  input (using LIKE). Return a formatted string with vendor name, country,
#  rating, and lead_days. Handle no-results case."
# =============================================================================

def get_approved_vendors(category: str) -> str:
    """
    Query NALCO's approved vendor database for a given procurement category.
    Returns a list of approved vendors with their ratings and lead times.
    
    Use this tool when you need to know which vendors are approved for a category such as 'Caustic Soda', 'Carbon Anodes', 'HFO', 'Aluminium Fluoride', etc.
    """
    conn = connect(DB_PATH)
    rows = conn.execute(
        "SELECT name, country, rating, lead_days FROM vendors WHERE category LIKE ?",
        (f"%{category}%",)
    ).fetchall()
    conn.close()

    if not rows:
        return f"No approved vendors found for category: {category}"

    # Format results as readable text the LLM can use in its reasoning
    lines = [
        f"Vendor: {r[0]} | Country: {r[1]} | Rating: {r[2]}/5.0 | Lead Time: {r[3]} days"
        for r in rows
    ]
    return f"Approved vendors for {category}:\n" + "\n".join(lines)


# =============================================================================
# TOOL 2: GET LAST PURCHASE PRICES
# PROMPT TO REGENERATE:
# "Write a tool function called get_last_prices(item: str) that
#  queries a SQLite table 'historical_pos' joined with 'vendors' for an item
#  matching the input (LIKE). Return the last 3 prices with dates and vendor names,
#  the average of those 3, and the current market benchmark from 'benchmark_rates'."
# =============================================================================

def get_last_prices(item: str) -> str:
    """
    Retrieve the last 3 purchase prices for a given item from NALCO's PO history, plus the current market benchmark rate.
    Use this tool to understand what NALCO has paid recently and how it compares to current market rates — essential context for setting RFQ evaluation criteria.
    """
    conn = connect(DB_PATH)

    # Get last 3 purchase prices
    rows = conn.execute(
        """SELECT h.item, h.price_inr, h.order_date, v.name
           FROM historical_pos h
           JOIN vendors v ON h.vendor_id = v.vendor_id
           WHERE h.item LIKE ?
           ORDER BY h.order_date DESC LIMIT 3""",
        (f"%{item}%",)
    ).fetchall()

    # Get current market benchmark
    bench = conn.execute(
        "SELECT market_rate_inr, unit FROM benchmark_rates WHERE item LIKE ?",
        (f"%{item}%",)
    ).fetchone()
    conn.close()

    if not rows:
        return f"No purchase history found for item: {item}"

    lines = [
        f"  {r[2]} | INR {r[1]:,.0f}/unit | Vendor: {r[3]}"
        for r in rows
    ]
    avg = sum(r[1] for r in rows) / len(rows)
    result = f"Last 3 purchase prices for {item}:\n" + "\n".join(lines)
    result += f"\n  Average of last 3: INR {avg:,.0f}"

    if bench:
        result += f"\n  Current market benchmark: INR {bench[0]:,.0f}/{bench[1]}"
        diff = avg - bench[0]
        result += f"\n  vs benchmark: {'INR {:,.0f} above'.format(diff) if diff > 0 else 'INR {:,.0f} below'.format(abs(diff))} market rate"

    return result


# =============================================================================
# TOOL 3: GENERATE RFQ DOCUMENT
# PROMPT TO REGENERATE:
# "Write a tool function called generate_rfq_docx(details: str)
#  that parses a pipe-delimited string 'item|qty|unit|delivery_date|vendor1,vendor2'.
#  Use python-docx to create a Word document with: title, RFQ number, date,
#  vendor list, item details table, and NALCO standard T&C bullets.
#  Save the file with a timestamp filename. Return the filename."
# =============================================================================

def generate_rfq_docx(details: str) -> str:
    """
    Generate a formatted Word RFQ document from procurement details.
    Input format: 'item|quantity|unit|delivery_date|vendor1,vendor2,vendor3'
    Example: 'Caustic Soda|500|MT|2024-07-15|Hindalco Chemicals,Indo Gulf Corp'
    
    Use this tool AFTER you have retrieved the approved vendors and last prices.
    """
    # Parse the pipe-delimited input string
    try:
        parts    = [p.strip() for p in details.split("|")]
        item     = parts[0] if len(parts) > 0 else "Item TBD"
        qty      = parts[1] if len(parts) > 1 else "TBD"
        unit     = parts[2] if len(parts) > 2 else "MT"
        del_date = parts[3] if len(parts) > 3 else str(date.today())
        vendors  = parts[4].split(",") if len(parts) > 4 else ["To be determined"]
    except Exception:
        item, qty, unit, del_date, vendors = details, "TBD", "MT", str(date.today()), ["TBD"]

    # Build the Word document using python-docx
    doc = Document()

    # Header section
    doc.add_heading("REQUEST FOR QUOTATION", level=0)
    doc.add_paragraph(f"RFQ No: NALCO/MAT/{date.today().strftime('%Y%m%d')}/AUTO-{item[:3].upper()}")
    doc.add_paragraph(f"Date: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph(f"To: {', '.join(vendors)}")
    doc.add_paragraph("")

    # Item details table
    doc.add_heading("Item Details", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Item / Material", "Quantity", "Unit", "Required Delivery Date"]):
        hdr[i].text = h

    row = table.add_row().cells
    row[0].text = item
    row[1].text = str(qty)
    row[2].text = unit
    row[3].text = del_date

    doc.add_paragraph("")

    # Terms and conditions
    doc.add_heading("Terms and Conditions", level=1)
    for tc in [
        "Delivery: As per schedule specified above. Delay penalty as per NALCO standard T&C.",
        "Payment: As per NALCO standard payment terms (45 days from delivery acceptance).",
        "Quality: Material must conform to NALCO technical specifications on file.",
        "Price: Quoted price must be firm for the full contract period.",
        "Validity: Quotation must remain valid for 45 days from date of submission.",
        "Reference: This RFQ is governed by NALCO Standard Procurement Terms & Conditions.",
    ]:
        doc.add_paragraph(tc, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Please submit your quote to: procurement@nalco.gov.in by the date specified above.")
    doc.add_paragraph("For queries: Contact NALCO Materials Management Department.")

    # Save the document with a unique filename
    filename = f"RFQ_{item.replace(' ', '_')}_{date.today().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(OUT_DIR, filename)
    doc.save(filepath)

    return f"RFQ document generated and saved: {filename}"

# =============================================================================
# DOWNLOAD BUTTON
# PROMPT TO REGENERATE:
# Write a Streamlit function `generate_docx_download_button()`
# that searches `OUT_DIR` for the most recently modified `RFQ_*.docx` file,
# loads it into a `BytesIO` object, displays an `st.download_button()`
# using the original filename and DOCX MIME type, then deletes the original file
# after creating the download button. Return only the function.
# =============================================================================
def generate_docx_download_button():
    '''
    Search for RFQ documents in the output directory and generate a streamlit button to download it.
    Remove the original file from the output directory after this.
    '''
    # Check if a new RFQ document was generated
    rfq_files = [
        f for f in os.listdir(OUT_DIR)
        if f.startswith("RFQ_") and f.endswith(".docx")
    ]

    if rfq_files:
        # Sort by modification time, get the most recent
        latest = max(rfq_files, key=lambda f: os.path.getmtime(os.path.join(OUT_DIR, f)))
        latest_filepath = os.path.join(OUT_DIR, latest)

        # Open docx file, convert to byte object and generate the download button
        with open(latest_filepath, "rb") as file:
            docx_buffer = BytesIO(file.read())

            st.download_button(
                label=f"Download Generated RFQ: {latest}",
                data=docx_buffer,
                file_name=latest,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        # Perform cleanup by removing the original docx file    
        os.remove(latest_filepath)
            
# =============================================================================
# STREAMLIT PAGE SETUP AND AGENT INITIALISATION
# PROMPT TO REGENERATE:
# "Set up a Streamlit page with title 'NALCO Procurement AI Agent'
#  and caption "Ollama + SQLite — 100% local, zero external API"
#  Register three tools in a list along with a dictionary of tool names.
# =============================================================================

st.set_page_config(page_title="UC-P04 | NALCO Procurement Agent", layout="wide")
st.title("UC-P04 | NALCO Procurement AI Agent")
st.caption("Ollama + SQLite — 100% local, zero external API")

# Register the three tools the agent can use
tools = [
    get_approved_vendors,
    get_last_prices,
    generate_rfq_docx,
    UC_P04_PO_email_tools.check_for_latest_quotation,
    UC_P04_PO_email_tools.generate_po,
    UC_P04_PO_email_tools.send_email
    ]

tool_names = {
            "get_approved_vendors" : get_approved_vendors,
            "get_last_prices" : get_last_prices,
            "generate_rfq_docx" : generate_rfq_docx,
            "check_for_latest_quotation" : UC_P04_PO_email_tools.check_for_latest_quotation,
            "generate_po" : UC_P04_PO_email_tools.generate_po,
            "send_email" : UC_P04_PO_email_tools.send_email
            }

# =============================================================================
# USER INTERFACE AND AGENT INVOCATION
# PROMPT TO REGENERATE:
# "Add two example instructions as markdown hints. Create a Streamlit text_input
#  for user instructions. Setup system prompt from 'https://smith.langchain.com/hub/hwchase17/react'.
#  Setup message structure with system prompt for ollama chat. Add a primary button 'Run Agent'. On click, invoke an
#  agent loop with ollama chat using the model 'qwen3.5', tools and the user input messages,
#  running till there are no tool calls to make.
#  Display the output, check for any .docx files created recently, and add a download button for the latest one.
#  Handle execution errors."
# =============================================================================
st.markdown("**Example instructions to try:**")
st.markdown(
    "- `Draft RFQ for 500 MT Caustic Soda, delivery Damanjodi by 15 July`\n"
    "- `Who are our approved vendors for Carbon Anodes and what did we pay last time?`\n"
    "- `Compare last 3 prices paid for HFO against market benchmark`"
)

st.markdown("""
**Watch the TERMINAL while the agent runs.** You will see:
- `Thought: I need to find vendors first...`
- `Action: get_approved_vendors('Caustic Soda')`
- `Observation: [results]`
- `Thought: Now I need prices...`

The agent is deciding the sequence — you never hard-coded this logic.
""")


user_input = st.text_input("Enter your procurement instruction:")

# Get system prompt from 'https://smith.langchain.com/hub/hwchase17/react' and setup the message structure
system_prompt = "input_variables=['agent_scratchpad', 'user_input', 'tool_names', 'tools']\ninput_types={}\npartial_variables={}\ntemplate='Answer the following questions as best you can. You have access to the following tools:\n\n{tools}\n\nUse the following format:\n\nQuestion: the input question you must answer\nThought: you should always think about what to do\nAction: the action to take, should be one of [{tool_names}]\nAction Input: the input to the action\nObservation: the result of the action\n...(Thought/Action/Action Input/Observation can repeat N times)\nThought: I now know the final answer\nFinal Answer: the final answer to the original input question\n\nBegin!\n\nQuestion: {user_input}\nThought: {agent_scratchpad}"
#messages = [{"role" : "system", "content" : system_prompt}]

# 1. Initialize the variable if it doesn't exist yet
if "messages" not in st.session_state:
    st.session_state.messages = [{"role" : "system", "content" : system_prompt}]

if st.button("Run Agent", type="primary") and user_input:

    st.session_state.messages.append({"role" : "user", "content" : user_input})
    with st.spinner("Agent working... (watch terminal for reasoning chain)"):
        try:
            
            while True:
                response = chat(model='qwen3.5:9b', messages=st.session_state.messages, tools=tools, think=True)
                st.session_state.messages.append(response.message)
                #messages.append(response.message)
                print("Thinking: ", response.message.thinking)
                print("Content: ", response.message.content)
                if response.message.tool_calls:
                    for tc in response.message.tool_calls:
                        if tc.function.name in tool_names:
                            print(f"Calling {tc.function.name} with arguments {tc.function.arguments}")
                            result = tool_names[tc.function.name](**tc.function.arguments)
                            print(f"Result: {result}")
                            # add the tool result to the messages
                            st.session_state.messages.append({'role': 'tool', 'tool_name': tc.function.name, 'content': str(result)})
                else:
                    # end the loop when there are no more tool calls
                    break
                # continue the loop with the updated messages

            st.markdown("---")
            st.markdown("### Agent Output")
            st.markdown(response.message.content)
            generate_docx_download_button()

        except Exception as e:
            st.error(f"Agent execution error: {e}")
            st.info("Check that Ollama is running (ollama serve) and the database exists (python UC_P04_setup_db.py).")